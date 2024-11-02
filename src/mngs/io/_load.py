#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Time-stamp: "2024-11-02 16:23:36 (ywatanabe)"
# File: ./mngs_repo/src/mngs/io/_load.py
#!/usr/bin/env python3

import json
import os
import pickle
import warnings

import h5py
import joblib
import mne
import numpy as np
import pandas as pd
import torch
import yaml
from PIL import Image


def load(lpath, show=False, verbose=False, **kwargs):
    """
    Load data from various file formats.

    This function supports loading data from multiple file formats including CSV, Excel, Numpy, Pickle, JSON, YAML, and more.

    Parameters:
    -----------
    lpath : str
        The path to the file to be loaded.
    show : bool, optional
        If True, display additional information during loading. Default is False.
    verbose : bool, optional
        If True, print verbose output during loading. Default is False.
    **kwargs : dict
        Additional keyword arguments to be passed to the specific loading function.

    Returns:
    --------
    object
        The loaded data object, which can be of various types depending on the input file format.

    Raises:
    -------
    ValueError
        If the file extension is not supported.

    Examples:
    ---------
    >>> data = load('data.csv')
    >>> image = load('image.png')
    >>> model = load('model.pth')
    """
    if lpath.startswith('f"'):
        try:
            lpath = eval(lpath)
        except:
            lpath = lpath.replace('f"', '')

    lpath = lpath.replace("/./", "/")

    try:
        extension = "." + lpath.split(".")[-1]

        # CSV
        if extension == ".csv":
            index_col = kwargs.get("index_col", 0)
            obj = pd.read_csv(lpath, **kwargs)
            obj = obj.loc[:, ~obj.columns.str.contains("^Unnamed")]
        # TSV
        elif extension == ".tsv":
            obj = pd.read_csv(lpath, sep="\t", **kwargs)
        # Excel
        elif extension in [".xls", ".xlsx", ".xlsm", ".xlsb"]:
            obj = pd.read_excel(lpath, **kwargs)
        # Parquet
        elif extension == ".parquet":
            obj = pd.read_parquet(lpath, **kwargs)
        # Numpy
        elif extension == ".npy":
            obj = np.load(lpath, allow_pickle=True, **kwargs)
        # Numpy NPZ
        elif extension == ".npz":
            obj = np.load(lpath)
            obj = dict(obj)
            obj = [v for v in obj.values()]
        # Pickle
        elif extension == ".pkl":
            with open(lpath, "rb") as l:
                obj = pickle.load(l, **kwargs)
        # Joblib
        elif extension == ".joblib":
            with open(lpath, "rb") as l:
                obj = joblib.load(l, **kwargs)
        # HDF5
        elif extension == ".hdf5":
            obj = {}
            with h5py.File(lpath, "r") as hf:
                for name in hf:
                    obj[name] = hf[name][:]
        # JSON
        elif extension == ".json":
            with open(lpath, "r") as f:
                obj = json.load(f)
        # Image
        elif extension in [".jpg", ".png", ".tiff", "tif"]:
            obj = Image.open(lpath)
        # YAML
        elif extension in [".yaml", ".yml"]:
            lower = kwargs.pop("lower", False)
            with open(lpath) as f:
                obj = yaml.safe_load(f, **kwargs)
            if lower:
                obj = {k.lower(): v for k, v in obj.items()}
        # Text
        elif extension in [".txt", ".log", ".event"]:
            with open(lpath, "r") as f:
                obj = f.read().splitlines()
                # obj = tqdm(f.read().splitlines(), desc=f"Reading {lpath}")
        # Markdown
        elif extension == ".md":
            obj = load_markdown(lpath, **kwargs)
        # PyTorch
        elif extension in [".pth", ".pt"]:
            obj = torch.load(lpath, **kwargs)
        # MATLAB
        elif extension == ".mat":
            from pymatreader import read_mat

            obj = read_mat(lpath, **kwargs)
        # XML
        elif extension == ".xml":
            from xml2dict import xml2dict

            obj = xml2dict(lpath, **kwargs)
        # CON
        elif extension == ".con":
            obj = mne.io.read_raw_fif(lpath, preload=True, **kwargs)
            obj = obj.to_data_frame()
            obj["samp_rate"] = obj.info["sfreq"]
        # CatBoost model
        elif extension == ".cbm":
            from catboost import CatBoostModel

            obj = CatBoostModel.load_model(lpath, **kwargs)

        # EEG data
        elif extension in [
            ".vhdr",
            ".vmrk",
            ".edf",
            ".bdf",
            ".gdf",
            ".cnt",
            ".egi",
            ".eeg",
            ".set",
        ]:
            obj = _load_eeg_data(lpath, **kwargs)

        elif extension == ".db":
            try:
                from ..db._BaseSQLiteDB import BaseSQLiteDB
                obj = BaseSQLiteDB(lpath)
            except Exception as e:
                raise ValueError(str(e))

        else:
            raise ValueError(f"Unsupported file extension: {extension}")

        return obj



    except Exception as e:
        print(f"Error loading file {lpath}: {str(e)}")
        raise


def _load_text(lpath):
    """
    Load text from a file.

    Parameters:
    -----------
    lpath : str
        The path to the text file to be loaded.

    Returns:
    --------
    str
        The content of the text file as a string.

    Raises:
    -------
    FileNotFoundError
        If the specified file does not exist.
    IOError
        If there's an error reading the file.
    """
    with open(lpath, "r") as f:
        return f.read()


def _load_eeg_data(filename, **kwargs):
    """
    Load EEG data based on file extension and associated files using MNE-Python.

    This function supports various EEG file formats including BrainVision, EDF, BDF, GDF, CNT, EGI, and SET.
    It also handles special cases for .eeg files (BrainVision and Nihon Koden).

    Parameters:
    -----------
    filename : str
        The path to the EEG file to be loaded.
    **kwargs : dict
        Additional keyword arguments to be passed to the specific MNE loading function.

    Returns:
    --------
    raw : mne.io.Raw
        The loaded raw EEG data.

    Raises:
    -------
    ValueError
        If the file extension is not supported.

    Notes:
    ------
    This function uses MNE-Python to load the EEG data. It automatically detects the file format
    based on the file extension and uses the appropriate MNE function to load the data.
    """
    # Get the file extension
    extension = filename.split(".")[-1]

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", RuntimeWarning)

        # Load the data based on the file extension
        if extension in ["vhdr", "vmrk"]:
            # Load BrainVision data
            raw = mne.io.read_raw_brainvision(filename, preload=True, **kwargs)
        elif extension == "edf":
            # Load European data format
            raw = mne.io.read_raw_edf(filename, preload=True, **kwargs)
        elif extension == "bdf":
            # Load BioSemi data format
            raw = mne.io.read_raw_bdf(filename, preload=True, **kwargs)
        elif extension == "gdf":
            # Load Gen data format
            raw = mne.io.read_raw_gdf(filename, preload=True, **kwargs)
        elif extension == "cnt":
            # Load Neuroscan CNT data
            raw = mne.io.read_raw_cnt(filename, preload=True, **kwargs)
        elif extension == "egi":
            # Load EGI simple binary data
            raw = mne.io.read_raw_egi(filename, preload=True, **kwargs)
        elif extension == "set":
            # ???
            raw = mne.io.read_raw(filename, preload=True, **kwargs)
        elif extension == "eeg":
            is_BrainVision = any(
                os.path.isfile(filename.replace(".eeg", ext))
                for ext in [".vhdr", ".vmrk"]
            )
            is_NihonKoden = any(
                os.path.isfile(filename.replace(".eeg", ext))
                for ext in [".21e", ".pnt", ".log"]
            )

            # Brain Vision
            if is_BrainVision:
                filename_v = filename.replace(".eeg", ".vhdr")
                raw = mne.io.read_raw_brainvision(
                    filename_v, preload=True, **kwargs
                )
            # Nihon Koden
            if is_NihonKoden:
                # raw = mne.io.read_raw_nihon(filename, preload=True, **kwargs)
                raw = mne.io.read_raw(filename, preload=True, **kwargs)
        else:
            raise ValueError(f"Unsupported file extension: {extension}")

        return raw


def load_markdown(lpath_md, style="plain_text"):
    """
    Load and convert Markdown content from a file.

    This function reads a Markdown file and converts it to either HTML or plain text format.

    Parameters:
    -----------
    lpath_md : str
        The path to the Markdown file to be loaded.
    style : str, optional
        The output style of the converted content.
        Options are "html" or "plain_text" (default).

    Returns:
    --------
    str
        The converted content of the Markdown file, either as HTML or plain text.

    Raises:
    -------
    FileNotFoundError
        If the specified file does not exist.
    IOError
        If there's an error reading the file.
    ValueError
        If an invalid style option is provided.

    Notes:
    ------
    This function uses the 'markdown' library to convert Markdown to HTML,
    and 'html2text' to convert HTML to plain text when necessary.
    """
    import html2text
    import markdown

    # Load Markdown content from a file
    with open(lpath_md, "r") as file:
        markdown_content = file.read()

    # Convert Markdown to HTML
    html_content = markdown.markdown(markdown_content)
    if style == "html":
        return html_content
    elif style == "plain_text":
        text_maker = html2text.HTML2Text()
        text_maker.ignore_links = True
        text_maker.bypass_tables = False
        plain_text = text_maker.handle(html_content)
        return plain_text
    else:
        raise ValueError(
            "Invalid style option. Choose 'html' or 'plain_text'."
        )


def _check_encoding(file_path):
    """
    Check the encoding of a given file.

    This function attempts to read the file with different encodings
    to determine the correct one.

    Parameters:
    -----------
    file_path : str
        The path to the file to check.

    Returns:
    --------
    str
        The detected encoding of the file.

    Raises:
    -------
    IOError
        If the file cannot be read or the encoding cannot be determined.
    """
    import chardet

    with open(file_path, "rb") as file:
        raw_data = file.read()

    result = chardet.detect(raw_data)
    return result["encoding"]


#     from chardet.universaldetector import UniversalDetector

#     detector = UniversalDetector()
#     with open(file_path, mode="rb") as f:
#         for binary in f:
#             detector.feed(binary)
#             if detector.done:
#                 break
#     detector.close()
#     enc = detector.result["encoding"]
#     return enc


def get_data_path_from_a_package(package_str, resource):
    """
    Get the path to a data file within a package.

    This function finds the path to a data file within a package's data directory.

    Parameters:
    -----------
    package_str : str
        The name of the package as a string.
    resource : str
        The name of the resource file within the package's data directory.

    Returns:
    --------
    str
        The full path to the resource file.

    Raises:
    -------
    ImportError
        If the specified package cannot be found.
    FileNotFoundError
        If the resource file does not exist in the package's data directory.
    """
    import importlib
    import os
    import sys

    spec = importlib.util.find_spec(package_str)
    if spec is None:
        raise ImportError(f"Package '{package_str}' not found")

    data_dir = os.path.join(spec.origin.split("src")[0], "data")
    resource_path = os.path.join(data_dir, resource)

    if not os.path.exists(resource_path):
        raise FileNotFoundError(
            f"Resource '{resource}' not found in package '{package_str}'"
        )

    return resource_path


def load_yaml_as_an_optuna_dict(fpath_yaml, trial):
    """
    Load a YAML file and convert it to an Optuna-compatible dictionary.

    This function reads a YAML file containing hyperparameter configurations
    and converts it to a dictionary suitable for use with Optuna trials.

    Parameters:
    -----------
    fpath_yaml : str
        The file path to the YAML configuration file.
    trial : optuna.trial.Trial
        The Optuna trial object to use for suggesting hyperparameters.

    Returns:
    --------
    dict
        A dictionary containing the hyperparameters with values suggested by Optuna.

    Raises:
    -------
    FileNotFoundError
        If the specified YAML file does not exist.
    ValueError
        If the YAML file contains invalid configuration for Optuna.
    """
    _d = load(fpath_yaml)

    for k, v in _d.items():
        dist = v["distribution"]

        if dist == "categorical":
            _d[k] = trial.suggest_categorical(k, v["values"])

        elif dist == "uniform":
            _d[k] = trial.suggest_int(k, float(v["min"]), float(v["max"]))

        elif dist == "loguniform":
            _d[k] = trial.suggest_loguniform(
                k, float(v["min"]), float(v["max"])
            )

        elif dist == "intloguniform":
            _d[k] = trial.suggest_int(
                k, float(v["min"]), float(v["max"]), log=True
            )

    return _d


def load_study_rdb(study_name, rdb_raw_bytes_url):
    """
    Load an Optuna study from a RDB (Relational Database) file.

    This function loads an Optuna study from a given RDB file URL.

    Parameters:
    -----------
    study_name : str
        The name of the Optuna study to load.
    rdb_raw_bytes_url : str
        The URL of the RDB file, typically in the format "sqlite:///*.db".

    Returns:
    --------
    optuna.study.Study
        The loaded Optuna study object.

    Raises:
    -------
    optuna.exceptions.StorageInvalidUsageError
        If there's an error loading the study from the storage.

    Example:
    --------
    >>> study = load_study_rdb(
    ...     study_name="YOUR_STUDY_NAME",
    ...     rdb_raw_bytes_url="sqlite:///path/to/your/study.db"
    ... )
    """
    import optuna

    storage = optuna.storages.RDBStorage(url=rdb_raw_bytes_url)
    study = optuna.load_study(study_name=study_name, storage=storage)
    print(f"Loaded: {rdb_raw_bytes_url}")
    return study




################################################################################
# dev
################################################################################
def _load_docx(lpath):
    """
    Load and extract text content from a .docx file.

    Parameters:
    -----------
    lpath : str
        The path to the .docx file.

    Returns:
    --------
    str
        The extracted text content from the .docx file.

    Raises:
    -------
    FileNotFoundError
        If the specified file does not exist.
    docx.opc.exceptions.PackageNotFoundError
        If the file is not a valid .docx file.
    """
    from docx import Document

    doc = Document(lpath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "".join(full_text)


#     doc = docx.Document(lpath)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return "\n".join(full_text)


# def _load_pdf(lpath):
#     reader = PyPDF2.PdfReader(lpath)
#     full_text = []
#     for page_num in range(len(reader.pages)):
#         page = reader.pages[page_num]
#         full_text.append(page.extract_text())
#     return "\n".join(full_text)


# def _load_latex(lpath):
#     return lpath.read().decode("utf-8")


# def _load_excel(lpath):
#     workbook = openpyxl.load_workbook(lpath)
#     all_text = []
#     for sheet in workbook:
#         for row in sheet.iter_rows(values_only=True):
#             all_text.append(
#                 " ".join(
#                     [str(cell) if cell is not None else "" for cell in row]
#                 )
#             )
#     return "\n".join(all_text)


# def _load_markdown(lpath):
#     md_text = StringIO(lpath.read().decode("utf-8"))
#     html = markdown.markdown(md_text.read())
#     return html


# def _load_textfile(lpath):
#     return lpath.read().decode("utf-8")


def load_markdown(lpath_md, style="plain_text"):
    """
    Load and convert a Markdown file to either HTML or plain text.

    Parameters:
    -----------
    lpath_md : str
        The path to the Markdown file.
    style : str, optional
        The output style, either "html" or "plain_text" (default).

    Returns:
    --------
    str
        The converted content of the Markdown file.
    """
    import html2text
    import markdown

    # Load Markdown content from a file
    with open(lpath_md, "r") as file:
        markdown_content = file.read()

    # Convert Markdown to HTML
    html_content = markdown.markdown(markdown_content)
    if style == "html":
        return html_content

    elif style == "plain_text":
        text_maker = html2text.HTML2Text()
        text_maker.ignore_links = True
        text_maker.bypass_tables = False
        plain_text = text_maker.handle(html_content)

        return plain_text


def load_yaml_as_an_optuna_dict(fpath_yaml, trial):
    """
    Load a YAML file and convert it to an Optuna-compatible dictionary.

    Parameters:
    -----------
    fpath_yaml : str
        The path to the YAML file.
    trial : optuna.trial.Trial
        The Optuna trial object.

    Returns:
    --------
    dict
        A dictionary with Optuna-compatible parameter suggestions.
    """
    _d = load(fpath_yaml)

    for k, v in _d.items():

        dist = v["distribution"]

        if dist == "categorical":
            _d[k] = trial.suggest_categorical(k, v["values"])

        elif dist == "uniform":
            _d[k] = trial.suggest_int(k, float(v["min"]), float(v["max"]))

        elif dist == "loguniform":
            _d[k] = trial.suggest_loguniform(
                k, float(v["min"]), float(v["max"])
            )

        elif dist == "intloguniform":
            _d[k] = trial.suggest_int(
                k, float(v["min"]), float(v["max"]), log=True
            )

    return _d


def load_study_rdb(study_name, rdb_raw_bytes_url):
    """
    Load an Optuna study from a RDB storage.

    Parameters:
    -----------
    study_name : str
        The name of the Optuna study.
    rdb_raw_bytes_url : str
        The URL of the RDB storage.

    Returns:
    --------
    optuna.study.Study
        The loaded Optuna study object.
    """
    import optuna

    # rdb_raw_bytes_url = "sqlite:////tmp/fake/ywatanabe/_MicroNN_WindowSize-1.0-sec_MaxEpochs_100_2021-1216-1844/optuna_study_test_file#0.db"
    storage = optuna.storages.RDBStorage(url=rdb_raw_bytes_url)
    study = optuna.load_study(study_name=study_name, storage=storage)
    print(f"\nLoaded: {rdb_raw_bytes_url}\n")
    return study


# EOF
