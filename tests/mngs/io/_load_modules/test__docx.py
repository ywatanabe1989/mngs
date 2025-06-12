#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Time-stamp: "2025-06-02 17:01:00 (ywatanabe)"
# File: ./mngs_repo/tests/mngs/io/_load_modules/test__docx.py

"""Comprehensive tests for DOCX file loading functionality."""

import os
import tempfile
import pytest
from unittest.mock import patch, MagicMock


class TestLoadDocx:
    """Test suite for _load_docx function"""
    
    def test_valid_extension_check(self):
        """Test that function validates .docx extension"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Test invalid extensions
        invalid_files = ["file.txt", "document.doc", "text.pdf", "data.xlsx"]
        
        for invalid_file in invalid_files:
            with pytest.raises(ValueError, match="File must have .docx extension"):
                _load_docx(invalid_file)
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_document_loading_and_text_extraction(self, mock_document_class):
        """Test that Document is loaded and text is extracted correctly"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock paragraphs
        mock_para1 = MagicMock()
        mock_para1.text = "This is the first paragraph."
        mock_para2 = MagicMock()
        mock_para2.text = "This is the second paragraph."
        mock_para3 = MagicMock()
        mock_para3.text = "Final paragraph here."
        
        # Create mock document
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_document_class.return_value = mock_doc
        
        # Call function
        result = _load_docx('test_document.docx')
        
        # Verify Document was called correctly
        mock_document_class.assert_called_once_with('test_document.docx')
        
        # Verify text extraction
        expected_text = "This is the first paragraph.This is the second paragraph.Final paragraph here."
        assert result == expected_text
        assert isinstance(result, str)
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_empty_document_handling(self, mock_document_class):
        """Test handling of empty DOCX documents"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock document with no paragraphs
        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('empty_document.docx')
        
        assert result == ""
        assert isinstance(result, str)
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_document_with_empty_paragraphs(self, mock_document_class):
        """Test handling of documents with empty paragraphs"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock paragraphs with some empty ones
        mock_para1 = MagicMock()
        mock_para1.text = "First paragraph."
        mock_para2 = MagicMock()
        mock_para2.text = ""  # Empty paragraph
        mock_para3 = MagicMock()
        mock_para3.text = "Third paragraph."
        mock_para4 = MagicMock()
        mock_para4.text = ""  # Another empty paragraph
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3, mock_para4]
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('mixed_document.docx')
        
        expected_text = "First paragraph.Third paragraph."
        assert result == expected_text
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_unicode_text_handling(self, mock_document_class):
        """Test handling of Unicode characters in document text"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock paragraphs with Unicode characters
        mock_para1 = MagicMock()
        mock_para1.text = "Héllo Wörld! 你好世界"
        mock_para2 = MagicMock()
        mock_para2.text = "Математика 🔬📊"
        mock_para3 = MagicMock()
        mock_para3.text = "العربية language test"
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('unicode_document.docx')
        
        expected_text = "Héllo Wörld! 你好世界Математика 🔬📊العربية language test"
        assert result == expected_text
        assert isinstance(result, str)
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_whitespace_and_special_characters(self, mock_document_class):
        """Test handling of whitespace and special characters"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock paragraphs with various whitespace scenarios
        mock_para1 = MagicMock()
        mock_para1.text = "  Leading and trailing spaces  "
        mock_para2 = MagicMock()
        mock_para2.text = "Tab\there\tand\tthere"
        mock_para3 = MagicMock()
        mock_para3.text = "Line\nbreaks\ninside"
        mock_para4 = MagicMock()
        mock_para4.text = "Special chars: !@#$%^&*()"
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3, mock_para4]
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('whitespace_document.docx')
        
        expected_text = "  Leading and trailing spaces  Tab\there\tand\tthereeLine\nbreaks\ninsideSpecial chars: !@#$%^&*()"
        assert result == expected_text
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_docx_exception_propagation(self, mock_document_class):
        """Test that python-docx exceptions are properly propagated"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Test FileNotFoundError
        mock_document_class.side_effect = FileNotFoundError("File not found")
        
        with pytest.raises(FileNotFoundError, match="File not found"):
            _load_docx('nonexistent.docx')
        
        # Test PackageNotFoundError (invalid DOCX file)
        from docx.opc.exceptions import PackageNotFoundError
        mock_document_class.side_effect = PackageNotFoundError("Invalid DOCX file")
        
        with pytest.raises(PackageNotFoundError, match="Invalid DOCX file"):
            _load_docx('invalid.docx')
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_large_document_handling(self, mock_document_class):
        """Test handling of large documents"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create a large number of mock paragraphs
        mock_paragraphs = []
        for i in range(1000):
            mock_para = MagicMock()
            mock_para.text = f"This is paragraph number {i}. "
            mock_paragraphs.append(mock_para)
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = mock_paragraphs
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('large_document.docx')
        
        # Verify it handles large documents without issues
        assert isinstance(result, str)
        assert len(result) > 20000  # Should be quite long
        assert "This is paragraph number 0. " in result
        assert "This is paragraph number 999. " in result
    
    def test_function_signature(self):
        """Test function signature and type annotations"""
        from mngs.io._load_modules._docx import _load_docx
        import inspect
        
        sig = inspect.signature(_load_docx)
        
        # Check parameters
        assert 'lpath' in sig.parameters
        assert 'kwargs' in sig.parameters or len(sig.parameters) >= 1
        
        # Check return annotation
        assert sig.return_annotation != inspect.Signature.empty
    
    def test_function_docstring(self):
        """Test that function has comprehensive docstring"""
        from mngs.io._load_modules._docx import _load_docx
        
        assert hasattr(_load_docx, '__doc__')
        assert _load_docx.__doc__ is not None
        docstring = _load_docx.__doc__
        
        # Check for key documentation elements
        assert 'Load and extract text content' in docstring
        assert 'Parameters' in docstring
        assert 'Returns' in docstring
        assert 'Raises' in docstring
        assert '.docx' in docstring
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_kwargs_ignored(self, mock_document_class):
        """Test that kwargs are ignored (not passed to Document)"""
        from mngs.io._load_modules._docx import _load_docx
        
        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_para.text = "Test text"
        mock_doc.paragraphs = [mock_para]
        mock_document_class.return_value = mock_doc
        
        # Call with various kwargs
        result = _load_docx('test.docx', verbose=True, encoding='utf-8', custom_param="value")
        
        # Verify Document was called only with the path
        mock_document_class.assert_called_once_with('test.docx')
        assert result == "Test text"
    
    def test_case_sensitive_extension_check(self):
        """Test case sensitivity of .docx extension"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Test case variations that should fail
        case_variations = ["file.DOCX", "file.Docx", "file.dOcX"]
        
        for variant in case_variations:
            with pytest.raises(ValueError, match="File must have .docx extension"):
                _load_docx(variant)
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_document_with_only_whitespace(self, mock_document_class):
        """Test handling of document with only whitespace paragraphs"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Create mock paragraphs with only whitespace
        mock_para1 = MagicMock()
        mock_para1.text = "   "
        mock_para2 = MagicMock()
        mock_para2.text = "\t\t"
        mock_para3 = MagicMock()
        mock_para3.text = "\n\n"
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('whitespace_only.docx')
        
        expected_text = "   \t\t\n\n"
        assert result == expected_text
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_absolute_and_relative_paths(self, mock_document_class):
        """Test handling of absolute and relative file paths"""
        from mngs.io._load_modules._docx import _load_docx
        
        mock_doc = MagicMock()
        mock_para = MagicMock()
        mock_para.text = "Test content"
        mock_doc.paragraphs = [mock_para]
        mock_document_class.return_value = mock_doc
        
        # Test absolute path
        abs_path = '/home/user/documents/report.docx'
        result = _load_docx(abs_path)
        mock_document_class.assert_called_with(abs_path)
        assert result == "Test content"
        
        # Reset mock and test relative path
        mock_document_class.reset_mock()
        rel_path = './data/document.docx'
        result = _load_docx(rel_path)
        mock_document_class.assert_called_with(rel_path)
        assert result == "Test content"
    
    @patch('mngs.io._load_modules._docx.Document')
    def test_complex_paragraph_text_scenarios(self, mock_document_class):
        """Test complex scenarios with paragraph text extraction"""
        from mngs.io._load_modules._docx import _load_docx
        
        # Test paragraphs with complex formatting (text should still be extracted)
        mock_para1 = MagicMock()
        mock_para1.text = "Bold and italic text mixed together"
        mock_para2 = MagicMock()
        mock_para2.text = "Text with hyperlinks and footnotes"
        mock_para3 = MagicMock()
        mock_para3.text = "Tables and lists converted to text"
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
        mock_document_class.return_value = mock_doc
        
        result = _load_docx('complex_formatting.docx')
        
        expected_text = "Bold and italic text mixed togetherText with hyperlinks and footnotesTables and lists converted to text"
        assert result == expected_text
    
    def test_import_error_handling(self):
        """Test that ImportError is handled when python-docx is not available"""
        with patch('builtins.__import__', side_effect=ImportError("No module named 'docx'")):
            from mngs.io._load_modules._docx import _load_docx
            
            # The import error should occur when Document is imported inside the function
            with pytest.raises(ImportError):
                _load_docx('test.docx')
    
    @pytest.mark.skipif(True, reason="Real file test - requires actual docx file creation")
    def test_real_docx_file_integration(self):
        """Integration test with real DOCX file (skipped by default)"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        
        from mngs.io._load_modules._docx import _load_docx
        
        # Create a real DOCX file
        doc = Document()
        doc.add_heading('Integration Test Document', 0)
        doc.add_paragraph('This is a real test paragraph.')
        doc.add_paragraph('Testing actual DOCX file loading.')
        
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            doc.save(f.name)
            temp_path = f.name
        
        try:
            content = _load_docx(temp_path)
            
            assert isinstance(content, str)
            assert 'Integration Test Document' in content
            assert 'real test paragraph' in content
            assert 'actual DOCX file loading' in content
            
        finally:
            os.unlink(temp_path)


if __name__ == "__main__":
    import os
    import pytest
<<<<<<< HEAD
    
    pytest.main([os.path.abspath(__file__), "-v"])
=======

    pytest.main([os.path.abspath(__file__)])

# --------------------------------------------------------------------------------
# Start of Source Code from: /data/gpfs/projects/punim2354/ywatanabe/mngs_repo/src/mngs/io/_load_modules/_docx.py
# --------------------------------------------------------------------------------
# #!/usr/bin/env python3
# # -*- coding: utf-8 -*-
# # Time-stamp: "2024-11-14 07:55:35 (ywatanabe)"
# # File: ./mngs_repo/src/mngs/io/_load_modules/_docx.py
# 
# from typing import Any
# 
# 
# def _load_docx(lpath: str, **kwargs) -> Any:
#     """
#     Load and extract text content from a .docx file.
# 
#     Parameters:
#     -----------
#     lpath : str
#         The path to the .docx file.
# 
#     Returns:
#     --------
#     str
#         The extracted text content from the .docx file.
# 
#     Raises:
#     -------
#     FileNotFoundError
#         If the specified file does not exist.
#     docx.opc.exceptions.PackageNotFoundError
#         If the file is not a valid .docx file.
#     """
#     if not lpath.endswith(".docx"):
#         raise ValueError("File must have .docx extension")
# 
#     from docx import Document
# 
#     doc = Document(lpath)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return "".join(full_text)
# 
# 
# # EOF

# --------------------------------------------------------------------------------
# End of Source Code from: /data/gpfs/projects/punim2354/ywatanabe/mngs_repo/src/mngs/io/_load_modules/_docx.py
# --------------------------------------------------------------------------------
>>>>>>> origin/main
